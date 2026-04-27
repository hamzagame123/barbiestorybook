from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
OUT_PATH = OUT_DIR / "barbie_storybook_stage3_asset_inventory_polished.docx"
SCREENSHOT_DIR = Path(r"C:\Users\HAMZA\Pictures\Screenshots")


@dataclass
class Asset:
    name: str
    preview: Path
    fmt: str
    size: str
    role: str
    note: str


PALETTE = {
    "bg": "13040D",
    "panel": "1C0813",
    "panel_alt": "250A18",
    "accent": "C93A74",
    "accent_dark": "6A1237",
    "text": "FFF5F8",
    "muted": "D8C4CC",
    "soft": "F3D6E2",
    "line": "5B2038",
}


def set_page_background(document: Document, color: str) -> None:
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color)
    document.element.insert(0, bg)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def write_text(paragraph, text: str, *, size=10, color=None, bold=False, font="Aptos"):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color or PALETTE["text"])
    return run


def title(document: Document, text: str, subtitle: str | None = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    write_text(p, text, size=24, bold=True, color=PALETTE["text"], font="Aptos Display")
    if subtitle:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(8)
        write_text(p, subtitle, size=11, color=PALETTE["soft"])


def section_header(document: Document, label: str, blurb: str) -> None:
    p = document.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(2)
    write_text(p, label, size=15, bold=True, color=PALETTE["soft"], font="Aptos Display")
    p = document.add_paragraph()
    p.space_after = Pt(8)
    write_text(p, blurb, size=10, color=PALETTE["muted"])


def callout(document: Document, heading: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, PALETTE["panel_alt"])
    margins(cell, 110, 130, 110, 130)
    p = cell.paragraphs[0]
    write_text(p, heading + "\n", size=12, bold=True, color=PALETTE["soft"], font="Aptos Display")
    write_text(p, body, size=10, color=PALETTE["text"])


def hero_panel(document: Document, image: Path, caption: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, PALETTE["panel"])
    margins(cell, 100, 100, 100, 100)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(6.2))
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_text(cap, caption, size=9, color=PALETTE["muted"])


def two_up(document: Document, left: tuple[Path, str], right: tuple[Path, str]) -> None:
    table = document.add_table(rows=1, cols=2)
    for idx, (img, cap) in enumerate((left, right)):
        cell = table.cell(0, idx)
        shade(cell, PALETTE["panel"])
        margins(cell, 70, 70, 90, 70)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Inches(2.95))
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        write_text(cp, cap, size=9, color=PALETTE["muted"])


def asset_cards(document: Document, heading: str, intro: str, assets: list[Asset]) -> None:
    section_header(document, heading, intro)
    for i in range(0, len(assets), 2):
        pair = assets[i:i + 2]
        table = document.add_table(rows=1, cols=2)
        if len(pair) == 1:
            table.cell(0, 1).text = ""
        for idx, asset in enumerate(pair):
            cell = table.cell(0, idx)
            shade(cell, PALETTE["panel"])
            margins(cell, 90, 90, 90, 90)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(asset.preview), width=Inches(2.3))

            name = cell.add_paragraph()
            name.alignment = WD_ALIGN_PARAGRAPH.LEFT
            write_text(name, asset.name, size=12, bold=True, color=PALETTE["soft"], font="Aptos Display")

            meta = cell.add_paragraph()
            write_text(meta, f"{asset.fmt}  |  {asset.size}", size=9, color=PALETTE["muted"], bold=True)

            role = cell.add_paragraph()
            write_text(role, asset.role, size=10, color=PALETTE["text"], bold=True)

            note = cell.add_paragraph()
            write_text(note, asset.note, size=9, color=PALETTE["muted"])


def bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        write_text(p, item, size=10, color=PALETTE["text"])


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_page_background(doc, PALETTE["bg"])

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10)

    title(doc, "BARBIE STORYBOOK", "Stage 3 - Building the Visual Language")
    callout(
        doc,
        "Visual Direction",
        "Barbie Storybook is an AR storytelling project built around glossy interface design, tabletop-scale 3D toys, dreamy panoramic backdrops, and a narrator character called Glimmer. This document summarizes the active visual system and the asset language currently used in the prototype.",
    )
    hero_panel(
        doc,
        SCREENSHOT_DIR / "Screenshot 2026-04-09 124038.png",
        "Main scan-state HUD with compact stage navigation and capture action.",
    )

    section_header(
        doc,
        "1. Interface Language",
        "The HUD uses a dark editorial shell, pink accent buttons, rounded glass surfaces, and a compact mobile AR layout. These screens show the active chrome and the narrator flow.",
    )
    two_up(
        doc,
        (SCREENSHOT_DIR / "Screenshot 2026-04-09 124100.png", "Start AR button treatment"),
        (SCREENSHOT_DIR / "Screenshot 2026-04-09 124055.png", "Library button treatment"),
    )
    two_up(
        doc,
        (SCREENSHOT_DIR / "Screenshot 2026-04-09 124104.png", "Book button treatment"),
        (SCREENSHOT_DIR / "Screenshot 2026-04-09 124043.png", "Expanded Glimmer helper panel"),
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_header(
        doc,
        "2. Brand Character: Glimmer",
        "Glimmer is the 2D narrator character used in the HUD. The character works as both a brand marker and an interaction cue for voice, text prompts, titles, beats, and captions.",
    )
    two_up(
        doc,
        (SCREENSHOT_DIR / "Screenshot 2026-04-09 124049.png", "Glimmer top-chrome badge"),
        (SCREENSHOT_DIR / "Screenshot 2026-04-09 124043.png", "Glimmer panel in use"),
    )
    bullet_list(
        doc,
        [
            "Type: 2D SVG UI character",
            "Role: narrator and story helper",
            "Visual language: pink-and-cream charm with gold trim and a soft kawaii face",
            "Technical value: lightweight, scalable, and easy to animate in the HUD",
        ],
    )

    dolls = [
        Asset("Classic Doll", ROOT / "src/assets/library/dolls/barbie-rigged.jpg", "GLB", "8.01 MB", "Primary doll preset", "Core Barbie figure for most staging scenarios."),
        Asset("Odile", ROOT / "src/assets/library/dolls/barbie-odile.jpg", "GLB", "7.02 MB", "Alternate doll preset", "Polished hero variant with strong toy-commercial styling."),
        Asset("Silkstone", ROOT / "src/assets/library/dolls/barbie-silkstone.jpg", "GLB", "6.67 MB", "Alternate doll preset", "Elegant silhouette with cleaner fashion proportions."),
        Asset("Deluxe Style", ROOT / "src/assets/library/dolls/barbie-deluxe-style.jpg", "GLB", "13.00 MB", "High-style variant", "Largest active doll asset and a more elaborate styling option."),
        Asset("Glimmer", SCREENSHOT_DIR / "Screenshot 2026-04-09 124049.png", "SVG", "5.54 KB", "2D UI character / narrator", "Glossy pink-and-cream star charm used as the story helper character."),
    ]
    asset_cards(
        doc,
        "3. Character System",
        "The character library includes four active doll presets plus Glimmer as the narrator character. Dolls are optimized for readable silhouettes in AR rather than extreme realism.",
        dolls,
    )

    backdrops = [
        Asset("Dreamhouse Pano", ROOT / "src/assets/worlds/barbie-dreamhouse-pano.jpg", "JPG", "0.64 MB", "Default backdrop", "Primary pano background anchoring the Barbie Dreamhouse setting."),
        Asset("Dollhouse Salon", ROOT / "src/assets/library/backgrounds/dollhouse-salon.jpg", "JPG", "4.57 MB", "Interior backdrop", "Warm dollhouse environment for fashion and social scenes."),
        Asset("Music Room", ROOT / "src/assets/library/backgrounds/dollhouse-music-room.jpg", "JPG", "6.09 MB", "Performance backdrop", "Largest active backdrop, used for music and stage-like moments."),
        Asset("Party Garden", ROOT / "src/assets/library/backgrounds/dollhouse-party-garden.jpg", "JPG", "4.41 MB", "Outdoor backdrop", "Celebration setting for brighter ending beats."),
    ]
    asset_cards(
        doc,
        "4. Backdrop System",
        "Backdrops are panoramic image environments designed to give quick narrative context without the weight of large explorable worlds.",
        backdrops,
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    accessories = [
        Asset("Toy Guitar", ROOT / "src/assets/library/accessories/toy-guitar.jpg", "GLB", "0.80 MB", "Music prop", "Strong silhouette for performance and stage moments."),
        Asset("Birthday Cake", ROOT / "src/assets/library/accessories/birthday-cake-imported.jpg", "GLB", "1.61 MB", "Celebration prop", "Useful for party scenes and ending-page captures."),
        Asset("Toy Puppy", ROOT / "src/assets/library/accessories/toy-puppy-imported.jpg", "GLB", "3.56 MB", "Companion prop", "Adds warmth and character to story beats."),
        Asset("Handbag", ROOT / "src/assets/library/accessories/handbag-imported.jpg", "GLB", "2.82 MB", "Fashion prop", "Simple accessory for styling and close-up beats."),
        Asset("Crown", ROOT / "src/assets/library/accessories/crown.jpg", "GLB", "0.70 MB", "Dress-up prop", "Small focal prop for close-up magical scenes."),
        Asset("Tawny Horse", ROOT / "src/assets/library/accessories/tawny-horse.jpg", "GLB", "1.09 MB", "Animal prop", "Larger companion object that broadens scene variety."),
        Asset("Barbie Movie Box", ROOT / "src/assets/library/accessories/barbie-box.jpg", "GLB", "0.52 MB", "Framing prop", "Photo-booth style prop useful for covers and playful staging."),
    ]
    asset_cards(
        doc,
        "5. Accessory Library",
        "Accessories carry most of the storytelling flexibility in the prototype. They are chosen for quick recognizability, readable silhouettes, and tabletop staging.",
        accessories,
    )

    section_header(
        doc,
        "6. Performance Summary",
        "The active library is designed for mobile AR. The strongest active performance risks come from the heavier doll models and the upscaled pano backdrops, while the accessory set remains relatively lightweight.",
    )
    bullet_list(
        doc,
        [
            "Active runtime set: 4 dolls, 4 backdrops, 7 accessories, and Glimmer as a 2D SVG narrator.",
            "Heaviest active doll: Deluxe Style at 13.00 MB.",
            "Largest active backdrop: Music Room at 6.09 MB.",
            "Accessory library remains comparatively lightweight and readable in AR.",
            "The overall visual language prioritizes glamorous readability and story clarity over photoreal detail.",
        ],
    )

    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    build()
