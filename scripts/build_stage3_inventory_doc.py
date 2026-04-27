from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DOCX = OUTPUT_DIR / "barbie_storybook_stage3_asset_inventory.docx"


SCREENSHOT_DIR = Path(r"C:\Users\HAMZA\Pictures\Screenshots")


@dataclass
class AssetRow:
    preview: Path | None
    name: str
    category: str
    fmt: str
    size_mb: str
    status: str
    notes: str


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_page_color(document: Document, color: str) -> None:
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color)
    document.element.insert(0, background)


def add_run_text(paragraph, text: str, size: int, color: str, bold=False, font="Courier New"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_title_block(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    add_run_text(p, "BARBIE STORYBOOK", 24, "FFF2F7", True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "Stage 3 - Building the Visual Language", 14, "F7BCD5", True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "Asset Inventory and Interface Language Overview", 11, "D7C1CB")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(
        p,
        "This document summarizes the active visual system, HUD language, narrator identity, and real-time asset library used in the current Barbie Storybook prototype.",
        10,
        "E8DDE2",
    )


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "2A0A1A")
    set_cell_margins(cell, 120, 140, 120, 140)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_text(p, title + "\n", 12, "FFD9E8", True)
    add_run_text(p, body, 10, "F1E5EA")


def add_section_heading(document: Document, title: str, subtitle: str | None = None) -> None:
    p = document.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    add_run_text(p, title, 16, "FFF1F7", True)
    if subtitle:
        p = document.add_paragraph()
        p.space_after = Pt(8)
        add_run_text(p, subtitle, 10, "D8C7CF")


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        add_run_text(p, item, 10, "EEE2E7")


def add_visual_grid(document: Document, screenshots: list[tuple[Path, str]]) -> None:
    table = document.add_table(rows=3, cols=2)
    table.autofit = False
    widths = [Inches(3.25), Inches(3.25)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            set_cell_margins(cell, 60, 60, 80, 60)
            set_cell_shading(cell, "180711")

    for idx, (image_path, caption) in enumerate(screenshots):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(3.0))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_text(cap, caption, 9, "D8C7CF", True)


def add_asset_table(document: Document, rows: list[AssetRow]) -> None:
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Preview", "Asset", "Category", "Format", "Size", "Status", "Design Notes"]
    header_row = table.rows[0]
    for idx, label in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, "551230")
        set_cell_margins(cell, 70, 70, 70, 70)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_text(p, label, 9, "FFF4F8", True)

    for row in rows:
        tr = table.add_row().cells
        for cell in tr:
            set_cell_shading(cell, "180711")
            set_cell_margins(cell, 60, 60, 60, 60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        preview_cell = tr[0]
        if row.preview and row.preview.exists():
            p = preview_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(row.preview), width=Inches(0.9))
        else:
            p = preview_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_text(p, "N/A", 8, "BFAEB6")

        values = [row.name, row.category, row.fmt, row.size_mb, row.status, row.notes]
        for cell, value in zip(tr[1:], values):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_text(p, value, 8, "F1E6EA", bold=False)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_page_color(document, "12030C")

    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Courier New"
    normal.font.size = Pt(10)

    add_title_block(document)
    add_callout(
        document,
        "Project Snapshot",
        "Barbie Storybook is an AR story-making experience built around placing dolls, accessories, and immersive backdrops into a scene, then capturing those moments into a book. The visual language combines glossy pink editorial UI, toy-scale 3D assets, and a narrator helper called Glimmer.",
    )

    add_section_heading(
        document,
        "1. Interface Language",
        "These screenshots capture the current HUD direction: compact chrome, rounded glass buttons, bold monospace typography, and a dedicated narrator presence.",
    )
    add_visual_grid(
        document,
        [
            (SCREENSHOT_DIR / "Screenshot 2026-04-09 124038.png", "Main HUD scan state"),
            (SCREENSHOT_DIR / "Screenshot 2026-04-09 124043.png", "Expanded Glimmer helper panel"),
            (SCREENSHOT_DIR / "Screenshot 2026-04-09 124049.png", "Glimmer top-chrome badge"),
            (SCREENSHOT_DIR / "Screenshot 2026-04-09 124055.png", "Library button treatment"),
            (SCREENSHOT_DIR / "Screenshot 2026-04-09 124100.png", "Start AR button treatment"),
            (SCREENSHOT_DIR / "Screenshot 2026-04-09 124104.png", "Book button treatment"),
        ],
    )

    add_section_heading(
        document,
        "2. Brand Character: Glimmer",
        "Glimmer is the narrator and story helper. In the interface, Glimmer appears as a small charm-like badge and expands into a voice-and-text prompt panel for titles, beats, and captions.",
    )
    add_bullets(
        document,
        [
            "Role: narrator helper for story titles, beats, and captions.",
            "Visual form: glossy pink charm icon with a soft kawaii expression.",
            "Runtime asset: lightweight SVG icon for clean HUD integration.",
            "Interaction: lives in the top chrome and opens into a dedicated helper panel.",
        ],
    )

    active_rows = [
        AssetRow(ROOT / "src/assets/library/dolls/barbie-rigged.jpg", "Classic Doll", "Doll", "GLB", "8.01 MB", "Active", "Primary doll preset for staging scenes."),
        AssetRow(ROOT / "src/assets/library/dolls/barbie-odile.jpg", "Odile", "Doll", "GLB", "7.02 MB", "Active", "Alternate hero doll with polished toy-commercial styling."),
        AssetRow(ROOT / "src/assets/library/dolls/barbie-silkstone.jpg", "Silkstone", "Doll", "GLB", "6.67 MB", "Active", "Elegant alternate doll with clean silhouette."),
        AssetRow(ROOT / "src/assets/library/dolls/barbie-deluxe-style.jpg", "Deluxe Style", "Doll", "GLB", "13.00 MB", "Active", "Largest active doll asset, used as a high-style variant."),
        AssetRow(ROOT / "src/assets/library/backgrounds/dollhouse-salon.jpg", "Dreamhouse Pano", "Backdrop", "JPG", "0.64 MB", "Active", "Default pano backdrop anchoring the Barbie Storybook setting."),
        AssetRow(ROOT / "src/assets/library/backgrounds/dollhouse-salon.jpg", "Dollhouse Salon", "Backdrop", "JPG", "4.57 MB", "Active", "Upscaled dollhouse interior with warm editorial lighting."),
        AssetRow(ROOT / "src/assets/library/backgrounds/dollhouse-music-room.jpg", "Music Room", "Backdrop", "JPG", "6.09 MB", "Active", "Largest active background, supports music and performance scenes."),
        AssetRow(ROOT / "src/assets/library/backgrounds/dollhouse-party-garden.jpg", "Party Garden", "Backdrop", "JPG", "4.41 MB", "Active", "Outdoor celebration backdrop for bright party scenes."),
        AssetRow(ROOT / "src/assets/library/accessories/toy-guitar.jpg", "Toy Guitar", "Accessory", "GLB", "0.80 MB", "Active", "Readable music prop that stages performance moments clearly."),
        AssetRow(ROOT / "src/assets/library/accessories/birthday-cake-imported.jpg", "Birthday Cake", "Accessory", "GLB", "1.61 MB", "Active", "Celebration prop for moment and ending beats."),
        AssetRow(ROOT / "src/assets/library/accessories/toy-puppy-imported.jpg", "Toy Puppy", "Accessory", "GLB", "3.56 MB", "Active", "Companion prop that adds warmth and story personality."),
        AssetRow(ROOT / "src/assets/library/accessories/handbag-imported.jpg", "Handbag", "Accessory", "GLB", "2.82 MB", "Active", "Fashion prop with simple readable silhouette."),
        AssetRow(ROOT / "src/assets/library/accessories/crown.jpg", "Crown", "Accessory", "GLB", "0.70 MB", "Active", "Small dress-up prop suited to close-up story moments."),
        AssetRow(ROOT / "src/assets/library/accessories/tawny-horse.jpg", "Tawny Horse", "Accessory", "GLB", "1.09 MB", "Active", "Large companion prop that broadens scene variety."),
        AssetRow(ROOT / "src/assets/library/accessories/barbie-box.jpg", "Barbie Movie Box", "Accessory", "GLB", "0.52 MB", "Active", "Photo-booth style frame for cover shots and playful staging."),
        AssetRow(SCREENSHOT_DIR / "Screenshot 2026-04-09 124049.png", "Glimmer Icon", "UI", "SVG", "0.01 MB", "Active", "Narrator UI icon; lightweight and scalable for the HUD."),
    ]

    add_section_heading(
        document,
        "3. Active Asset Inventory",
        "This table lists the assets currently wired into the app flow. Preview images are included where available to show the visual language rather than just the filenames.",
    )
    add_asset_table(document, active_rows)

    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_section_heading(
        document,
        "4. Legacy and Reference Assets",
        "These assets still exist in the repository, but they are not part of the main active user flow. They are included here for technical completeness and future cleanup planning.",
    )
    legacy_rows = [
        AssetRow(None, "Test Barbie", "Doll", "GLB", "33.41 MB", "Legacy", "Very large test asset and not part of the active preset library."),
        AssetRow(None, "Park Playset", "Stage Prop", "FBX", "1.71 MB", "Legacy", "Older stage-prop asset from the earlier scene model."),
        AssetRow(None, "Dream House Model", "Stage Prop", "FBX", "19.98 MB", "Legacy", "Large legacy Dream House scene asset."),
    ]
    add_asset_table(document, legacy_rows)

    add_section_heading(
        document,
        "5. Performance Notes",
        "The current prototype is optimized around readable toy silhouettes and compact AR interactions instead of photorealism. Active assets are largely lightweight GLBs, while the remaining performance risks come from oversized legacy dolls and large upscaled pano backdrops.",
    )
    add_bullets(
        document,
        [
            "Active runtime set: 4 dolls, 4 backdrops, 7 accessories, and 1 narrator icon.",
            "Heaviest active 3D asset: Deluxe Style doll at 13.00 MB.",
            "Largest active backdrop: Music Room at 6.09 MB.",
            "Largest legacy risk: Test Barbie at 33.41 MB.",
            "The current visual system favors glossy, high-contrast readability over dense realism to support mobile AR.",
        ],
    )

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
